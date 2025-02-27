# Monkey-patch inspect to add getargspec for compatibility with phi
import inspect
if not hasattr(inspect, 'getargspec'):
    inspect.getargspec = inspect.getfullargspec

import streamlit as st
import os
import base64
from io import BytesIO
from docx import Document
from dotenv import load_dotenv
import requests
from pydantic import BaseModel

# Load environment variables (make sure .env contains GROQ_API_KEY)
load_dotenv()

# --- PHIDATA Agent Imports ---
from phi.agent import Agent, RunResponse
# Import the expected base model class from phi (adjust the import if needed)
from phi.model.base import Model as PhiModel

# Define a custom GroqLLM wrapper class as a subclass of the expected PhiModel
class GroqLLM(PhiModel):
    id: str
    api_key: str
    temperature: float = 0.1
    max_tokens: int = 8000

    def generate(self, prompt: str) -> str:
        # Use the correct endpoint as per Groq's API docs.
        url = "https://api.groq.com/openai/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        # Include the 'model' property along with a 'messages' array.
        data = {
            "model": self.id,  # e.g., "llama3-8b-8192"
            "messages": [{"role": "user", "content": prompt}],
            "temperature": self.temperature,
            "max_tokens": self.max_tokens
        }
        response = requests.post(url, headers=headers, json=data)
        if response.status_code == 200:
            json_data = response.json()
            try:
                # Extract the assistant's reply from the response
                return json_data["choices"][0]["message"]["content"]
            except Exception as e:
                return f"Error extracting response: {e}"
        else:
            return f"Error: {response.status_code} {response.text}"

    def response(self, messages: list) -> RunResponse:
        """
        Combine the incoming messages into a prompt, generate a response,
        and wrap it in a RunResponse object.
        """
        prompt = "\n".join([getattr(msg, "content", str(msg)) for msg in messages])
        generated_text = self.generate(prompt)
        return RunResponse(content=generated_text)

# Import the tools from Phidata
from phi.tools.duckduckgo import DuckDuckGo
from phi.tools.newspaper4k import Newspaper4k

# Initialize the Groq LLM model (using our custom Pydantic-based PhiModel)
groq_model = GroqLLM(id="llama3-8b-8192", api_key=os.getenv("GROQ_API_KEY"))

# --- Define the Medical Report Analyzer Agent ---
medical_report_agent = Agent(
    model=groq_model,
    description=(
        "You are a medical diagnosis expert who handles PHI data securely. "
        "Given a patient's medical report, name, age, and self-reported symptoms, analyze the report, "
        "extract key insights, and provide a preliminary diagnosis."
    ),
    instructions=[
        "Read and analyze the medical report.",
        "Include the patient's name and age in your analysis.",
        "Consider the patient's reported symptoms.",
        "Extract key insights and produce a preliminary diagnosis.",
        "Ensure your approach is HIPAA-compliant."
    ],
    markdown=True,
    show_tool_calls=True,
    add_datetime_to_instructions=True,
    # debug_mode=True,  # Uncomment for debugging details.
)

# --- Define the Web Information Collector Agent ---
web_info_agent = Agent(
    model=groq_model,
    tools=[DuckDuckGo(), Newspaper4k()],
    description=(
        "You are an AI that gathers up-to-date medical research and online data regarding patient symptoms. "
        "Search the web for relevant medical information, extract key findings, and summarize possible conditions and treatments."
    ),
    instructions=[
        "For the given symptoms, perform a web search using the available tools.",
        "Extract key points from the top search results.",
        "Summarize your findings in a concise manner."
    ],
    markdown=True,
    show_tool_calls=True,
    add_datetime_to_instructions=True,
    # debug_mode=True,
)

# --- Helper Functions ---
def generate_docx(report_text):
    """Generate a Word document from the diagnosis report text."""
    doc = Document()
    doc.add_heading("Full Diagnosis Report", 0)
    doc.add_paragraph(report_text)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def get_download_link(bio, filename):
    """Return a download link for the generated Word document."""
    b64 = base64.b64encode(bio.read()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download Diagnosis Report</a>'

# --- Streamlit UI ---
st.set_page_config(layout="wide")
st.title("PHI Compliant AI Health Agent using Phidata and Groq")

# Patient information inputs
patient_name = st.text_input("Enter your name", "John Doe")
patient_age = st.number_input("Enter your age", min_value=0, max_value=120, value=30, step=1)

# User input: symptoms and medical report file
symptoms = st.text_area("Enter your symptoms", "e.g., fever, cough, headache")
uploaded_file = st.file_uploader("Upload your medical report (TXT or DOCX)", type=["txt", "docx"])

if st.button("Generate Diagnosis Report"):
    with st.spinner("Generating diagnosis report..."):
        # Read the uploaded medical report file
        report_text = ""
        if uploaded_file is not None:
            if uploaded_file.type == "text/plain":
                report_text = uploaded_file.read().decode("utf-8")
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                from docx import Document as DocxDocument
                doc = DocxDocument(uploaded_file)
                paragraphs = [para.text for para in doc.paragraphs]
                report_text = "\n".join(paragraphs)
            else:
                report_text = "Unsupported file format provided."
        
        # Build prompts for each agent including patient info
        analysis_prompt = (
            f"Patient Name: {patient_name}\n"
            f"Patient Age: {patient_age}\n"
            f"Medical Report:\n{report_text}\n\n"
            f"Symptoms:\n{symptoms}"
        )
        web_prompt = (
            f"Patient Name: {patient_name}\n"
            f"Patient Age: {patient_age}\n"
            f"Symptoms:\n{symptoms}"
        )
        
        # Run the Medical Report Analyzer Agent
        analysis_response = medical_report_agent.run(analysis_prompt)
        
        # Run the Web Information Collector Agent
        web_response = web_info_agent.run(web_prompt)
        
        # Combine both agents' outputs into a full diagnosis report
        full_report = (
            f"--- Medical Report Analysis ---\n{analysis_response.content}\n\n"
            f"--- Online Medical Information ---\n{web_response.content}"
        )
        
        # Display the full diagnosis report in Streamlit
        st.write(full_report)
        
        # Generate a downloadable Word document for the report
        docx_file = generate_docx(full_report)
        download_link = get_download_link(docx_file, "Diagnosis_Report.docx")
        st.markdown(download_link, unsafe_allow_html=True)
