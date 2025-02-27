# main.py

# --- Compatibility patch for inspect (if needed)
import sys
import inspect
if not hasattr(inspect, 'getargspec'):
    inspect.getargspec = inspect.getfullargspec

import streamlit as st
import os
import base64
from io import BytesIO
from docx import Document
from dotenv import load_dotenv
import requests
from pydantic import BaseModel
from PIL import Image

# Load environment variables (ensure .env contains GROQ_API_KEY)
load_dotenv()

# --- PHIDATA Agent Imports ---
from phi.agent import Agent, RunResponse
# Import the expected base model class from phi (adjust the import if needed)
from phi.model.base import Model as PhiModel

# --- Custom GroqLLM Wrapper ---
class GroqLLM(PhiModel):
    id: str
    api_key: str
    temperature: float = 0.1
    max_tokens: int = 8000

    def generate(self, prompt: str) -> str:
        url = "https://api.groq.com/openai/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        data = {
            "model": self.id,  # e.g., "llama3-8b-8192"
            "messages": [{"role": "user", "content": prompt}],
            "temperature": self.temperature,
            "max_tokens": self.max_tokens
        }
        response = requests.post(url, headers=headers, json=data)
        if response.status_code == 200:
            json_data = response.json()
            try:
                return json_data["choices"][0]["message"]["content"]
            except Exception as e:
                return f"Error extracting response: {e}"
        else:
            return f"Error: {response.status_code} {response.text}"

    def response(self, messages: list) -> RunResponse:
        prompt = "\n".join([getattr(msg, "content", str(msg)) for msg in messages])
        generated_text = self.generate(prompt)
        return RunResponse(content=generated_text)

# Initialize the GroqLLM model using the GROQ_API_KEY from your .env file.
groq_model = GroqLLM(id="llama3-8b-8192", api_key=os.getenv("GROQ_API_KEY"))

# --- Define the Fitness Plan Agent ---
fitness_plan_agent = Agent(
    model=groq_model,
    description=(
        "You are a fitness and nutrition expert. Given a user's basic details, a photo analysis result, "
        "and an input prompt, calculate the Body Measurement Level (BML) and generate a personalized exercise plan and diet plan."
    ),
    instructions=[
        "Analyze the provided photo information (including a calculated BML value).",
        "Incorporate the user's additional prompt into your analysis.",
        "Generate a detailed, practical exercise plan and diet plan.",
        "Ensure recommendations are safe and tailored to the user's profile."
    ],
    markdown=True,
    show_tool_calls=True,
    add_datetime_to_instructions=True,
)

# --- Helper Functions ---

def generate_docx(report_text):
    """Generate a Word document from the report text."""
    doc = Document()
    doc.add_heading("Full Report", 0)
    doc.add_paragraph(report_text)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def get_download_link(bio, filename):
    """Return a download link for the generated Word document."""
    b64 = base64.b64encode(bio.read()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download Report</a>'

def calculate_bml(image: Image.Image) -> float:
    """
    Dummy function to calculate Body Measurement Level (BML) based on image properties.
    Replace this with proper image analysis in a real application.
    For now, it uses the width-to-height ratio scaled arbitrarily.
    """
    width, height = image.size
    if height == 0:
        return 0.0
    bml = (width / height) * 10  # arbitrary scaling factor
    return round(bml, 1)

# --- Streamlit UI ---

st.set_page_config(layout="wide")
st.title("Exercise & Diet Plan Generator")



# Patient information inputs
patient_name = st.text_input("Enter your name", "John Doe")
patient_age = st.number_input("Enter your age", min_value=0, max_value=120, value=30, step=1)

# --- Fitness Plan Section ---
st.header("Exercise & Diet Plan Generator")
uploaded_photo = st.file_uploader("Upload your photo", type=["jpg", "jpeg", "png"])
additional_prompt = st.text_area("Enter any additional details or fitness goals", "e.g., I want to build muscle and lose fat")

if st.button("Generate Fitness Report"):
    if uploaded_photo is not None:
        with st.spinner("Analyzing photo and generating fitness plan..."):
            try:
                image = Image.open(uploaded_photo)
            except Exception as e:
                st.error(f"Error opening image: {e}")
                image = None

            if image:
                # Calculate a dummy BML value from the image
                bml_value = calculate_bml(image)
                st.write(f"Calculated BML: {bml_value}")

                # Build prompt for the fitness plan agent
                fitness_prompt = (
                    f"Patient Name: {patient_name}\n"
                    f"Patient Age: {patient_age}\n"
                    f"Calculated BML (Body Measurement Level): {bml_value}\n"
                    f"Additional Fitness Goals/Prompt:\n{additional_prompt}"
                )

                fitness_response = fitness_plan_agent.run(fitness_prompt)
                st.write("### Your Personalized Exercise & Diet Plan")
                st.write(fitness_response.content)

                docx_file = generate_docx(fitness_response.content)
                download_link = get_download_link(docx_file, "Fitness_Report.docx")
                st.markdown(download_link, unsafe_allow_html=True)
    else:
        st.warning("Please upload a photo to generate the fitness plan.")
